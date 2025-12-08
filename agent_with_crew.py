
import os
import zipfile
import json
from pathlib import Path
from typing import Dict, Any, Optional
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from crewai import Agent, Task, Crew, Process
from crewai import LLM

# ====================================================
# 2. CONNECT TO YOUR AI FOUNDRY DEPLOYED MODEL
# ====================================================

# Azure AI Foundry Configuration
# For CrewAI with Azure, the endpoint needs the full inference URL
AZURE_ENDPOINT = "https://canaifoundry.cognitiveservices.azure.com/openai/deployments/gpt-4o-mini"
AZURE_API_KEY = "6uzWpypcLFsKXaDFGXiICvVt7GKcJxHNmFEKOaKuAleRdlY0iB1mJQQJ99BJAC77bzfXJ3w3AAAAACOGYs5G"
AZURE_DEPLOYMENT = "gpt-4o-mini"
AZURE_API_VERSION = "2024-12-01-preview"

# Set environment variables required by CrewAI for Azure
os.environ["AZURE_API_KEY"] = AZURE_API_KEY
os.environ["AZURE_ENDPOINT"] = AZURE_ENDPOINT
os.environ["AZURE_API_VERSION"] = AZURE_API_VERSION
os.environ["OPENAI_API_KEY"] = AZURE_API_KEY  # For embeddings

# Create LLM instance using CrewAI's LLM class for Azure
llm = LLM(
    model=f"azure/{AZURE_DEPLOYMENT}",
    api_key=AZURE_API_KEY,
    base_url=AZURE_ENDPOINT,
    api_version=AZURE_API_VERSION,
    temperature=0.3
)

# ====================================================
# 3. TOOLS
# ====================================================

from crewai.tools import BaseTool
from pydantic import Field

class CodebaseAnalyzerTool(BaseTool):
    name: str = "analyze_codebase"
    description: str = "Deep analysis of uploaded codebase (folder or zip). Extracts language, files, endpoints, auth methods, and security issues."
    
    def _run(self, codebase_path: str) -> Dict[str, Any]:
        """Execute the codebase analysis"""
        result = {
            "language": "unknown",
            "files": [],
            "endpoints": [],
            "auth_methods": [],
            "databases": [],
            "security_issues": [],
            "architecture": "Not detected"
        }

        extract_dir = "temp_codebase"
        if codebase_path.endswith('.zip'):
            with zipfile.ZipFile(codebase_path, 'r') as z:
                z.extractall(extract_dir)
            codebase_path = extract_dir

        for root, _, files in os.walk(codebase_path):
            for file in files:
                if file.endswith(('.py', '.js', '.ts', '.java', '.go', '.cs')):
                    path = os.path.join(root, file)
                    result["files"].append(path)
                    try:
                        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                            content = f.read().lower()
                            if any(x in content for x in ['fastapi', 'flask', '@app.route', 'router.']):
                                result["language"] = "Python Web"
                                result["endpoints"].extend([l.strip() for l in open(path) if 'get(' in l or 'post(' in l])
                            if 'express' in content:
                                result["language"] = "Node.js Express"
                            if any(x in content for x in ['jwt', 'oauth', 'passport', 'auth0']):
                                result["auth_methods"].append("OAuth/JWT detected")
                            if 'password' in content and ('=' in content or 'hardcoded' in content):
                                result["security_issues"].append(f"Possible secret in {file}")
                    except: pass

        result["architecture"] = f"{len(result['files'])} files | {result['language']} | {len(result['endpoints'])} endpoints"
        return result

# Initialize the tool
analyze_codebase = CodebaseAnalyzerTool()

class DocumentGeneratorTool(BaseTool):
    name: str = "generate_documents"
    description: str = "Generate PDF, DOCX, and Markdown files from documentation content"
    
    def _run(self, content: str, project_name: str = "Documentation_Package") -> str:
        """Generate actual document files"""
        try:
            # Save Markdown
            md_path = f"{project_name}.md"
            with open(md_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            # Generate DOCX
            docx_path = f"{project_name}.docx"
            doc = Document()
            
            # Add title
            title = doc.add_heading('Documentation Package', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Split content by lines and add to document
            lines = content.split('\n')
            for line in lines:
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.strip():
                    doc.add_paragraph(line)
            
            doc.save(docx_path)
            
            # Note: PDF generation requires weasyprint or reportlab
            # For now, we'll skip PDF to avoid dependency issues
            
            return f"Files generated successfully:\n- {md_path}\n- {docx_path}"
        except Exception as e:
            return f" Error generating files: {str(e)}"

generate_documents = DocumentGeneratorTool()

# ====================================================
# 4. THE 4 CORE AGENTS
# ====================================================

orchestrator = Agent(
    role="Documentation Director",
    goal="Review and validate all documentation for completeness and quality",
    backstory="15+ years leading enterprise architecture and documentation teams at Fortune 500s",
    llm=llm,
    allow_delegation=False,
    verbose=True
)

code_analyst = Agent(
    role="Senior Code & System Intelligence Engineer",
    goal="Deeply understand any codebase and extract architecture, risks, flows",
    backstory="Principal Engineer who reverse-engineers systems in minutes",
    tools=[analyze_codebase],
    llm=llm,
    verbose=True
)

doc_specialist = Agent(
    role="Enterprise Documentation Architect",
    goal="Write flawless BRD, FRD, NFRD, Security, Architecture docs",
    backstory="Authored 1000+ audit-ready docs for banks, fintechs, governments",
    llm=llm,
    verbose=True
)

formatter = Agent(
    role="Professional Document Publisher",
    goal="Convert content into beautiful DOCX and MD files with professional formatting",
    backstory="Designs documents that impress CTOs and compliance teams",
    tools=[generate_documents],
    llm=llm,
    verbose=True
)

# ====================================================
# 5. TASKS (Dynamic based on input)
# ====================================================

def create_docagent_crew(user_prompt: str, codebase_path: Optional[str] = None):
    inputs = {
        "user_prompt": user_prompt,
        "has_codebase": "Yes" if codebase_path else "No"
    }

    task1 = Task(
        description=f"Analyze the codebase at {codebase_path or 'N/A'} and return full JSON summary",
        expected_output="Structured JSON with architecture, endpoints, risks, auth",
        agent=code_analyst,
        async_execution=False
    )

    task2 = Task(
        description=f"""Using the code analysis + user prompt: '{user_prompt}'
        Generate complete professional documentation including:
        - Business Requirements Document (BRD)
        - Functional Requirements Document (FRD)
        - Non-Functional Requirements (NFRD)
        - Security & Compliance Document
        - Architecture Overview
        Use enterprise standards.""",
        expected_output="Full Markdown documentation with clear sections",
        agent=doc_specialist,
        context=[task1]
    )

    task3 = Task(
        description="""Take the final Markdown documentation and use the generate_documents tool to create actual files.
        Pass the complete markdown content to the tool to generate Documentation_Package.docx and Documentation_Package.md files.
        Return the confirmation message from the tool.""",
        expected_output="Confirmation message that DOCX and MD files have been generated successfully",
        agent=formatter,
        context=[task2]
    )

    task4 = Task(
        description="""Review the generated documentation and file generation results. 
        Provide a brief quality assessment covering:
        1. Completeness of all required sections
        2. Professional quality and clarity
        3. File generation success
        Keep the review concise and action-oriented.""",
        expected_output="Brief quality assessment and confirmation that all deliverables are complete",
        agent=orchestrator,
        context=[task1, task2, task3]
    )

    crew = Crew(
        agents=[orchestrator, code_analyst, doc_specialist, formatter],
        tasks=[task1, task2, task3, task4],
        process=Process.sequential,  # Change to hierarchical for advanced mode
        verbose=True,
        memory=False,  # Disable memory to avoid embedding issues
        cache=False     # Disable cache to avoid issues
    )

    return crew

# ====================================================
# 6. RUN THE AGENT
# ====================================================

def run_docagent(user_prompt: str, codebase_zip_path: Optional[str] = None):
    print("Starting DocAgent with your AI Foundry model...")
    crew = create_docagent_crew(user_prompt, codebase_zip_path)
    result = crew.kickoff()
    
    print("\n" + "="*60)
    print("DOCAGENT COMPLETE!")
    print("="*60)
    print("\nCheck your workspace for generated files:")
    print("   • Documentation_Package.docx")
    print("   • Documentation_Package.md")
    print("="*60)
    
    return result

# ====================================================
# 7. USAGE EXAMPLE
# ====================================================

if __name__ == "__main__":
    # Example without codebase analysis
    result = run_docagent(
        user_prompt="Generate full documentation for instagram clone application",
        codebase_zip_path=None  # No codebase - will generate based on prompt only
    )
    print(result)